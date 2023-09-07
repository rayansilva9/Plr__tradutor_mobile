import os
import time
from googletrans import Translator
from docx import Document

list = os.listdir("//storage//emulated//0//Documents//Pydroid3//PLR//")


tradutor = Translator(service_urls=["translate.google.com"])


arquivos = []
items_docx = []
items_txt = []

local_de_salvamento = "//storage//emulated//0//Documents//Pydroid3//TRADUZIDOS"

local_de_arquivos = "//storage//emulated//0//Documents//Pydroid3//PLR//"

traduzir_para = ["pt","en","fr"]



def dividir_texto(string, tamanho):
	return [string[i:i+tamanho] for i in range(0, len(string), tamanho)]
	
	
	
def obterArquivos():
    arquivos.clear()
    items_docx.clear()
    items_txt.clear()
    for i in list:
        arquivos.append(i)
        
    for plr in arquivos:
        nome, extensao = os.path.splitext(plr)
        
        if extensao == '.txt':
                items_txt.append({'nome':os.path.basename(nome), 'path':plr})
                
        if extensao == '.docx' :
                items_docx.append({'nome':os.path.basename(nome), 'path':plr})
                
                
def traduzir_txt():
    for plr in items_txt:
        for lang in traduzir_para:
            if os.path.exists(f"{local_de_salvamento}//{lang}"):
                pass
            else:
		os.makedirs(f"{local_de_salvamento}//{lang}")
            texto_completo1 = []
            texto_completo2 = ""
        
            nome_trad = tradutor.translate(plr["nome"].replace("_"," "), src="auto", dest=lang).text
        
            with open(local_de_arquivos + "/" + plr["path"], "r", encoding="cp1252") as arquivo:
                partes = dividir_texto(arquivo.read(), 4999)
                for parte in partes:
                    texto_completo1.append(tradutor.translate(parte, src="auto", dest=lang).text)
            
                texto_completo2 = "\n".join(texto_completo1)
                arquivo.close()
                
            
            
            with open(local_de_salvamento + "//" + lang + "//" + nome_trad + ".txt", "w") as arquivo:
                arquivo.write(texto_completo2)
                arquivo.close()
                print(nome_trad + "traduzido")
                
                
def traduzir_docx():
    
    for plr in items_docx:
        for lang in traduzir_para:
            if os.path.exists(f"{local_de_salvamento}//{lang}"):
                pass
            else:
                os.makedirs(f"{local_de_salvamento}/{lang}")
            
            texto1 = []
            texto2 = ""
            texto_completo1 = []
            texto_completo2 = ""
            documentTraduzido = Document()
            
            nome_trad = tradutor.translate(plr["nome"], src="auto", dest=lang).text
            
            doc = Document(local_de_arquivos  + "/" + plr["path"])
            doc.save(f"{local_de_salvamento}//{lang}//{nome_trad + '.docx'}")
            
            documentDocx = Document(f"{local_de_salvamento}//{lang}//{nome_trad + '.docx'}")
            
            
            for paragraph in documentDocx.paragraphs:
               for run in paragraph.runs:
                   try:
                       text = run.text
                       texto_traduzido = tradutor.translate(text, src="auto", dest=lang).text
                       run.text = texto_traduzido
                   except TypeError:
                       pass
                   except IndexError:
                       pass
                   documentDocx.save(local_de_salvamento + "/" + lang + "/" + nome_trad + ".docx")
            print(f"{plr['nome']}: traduzido")
               
            
obterArquivos()
traduzir_txt()
traduzir_docx()

